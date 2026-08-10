import io
import re
from typing import Any, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib.pagesizes import LETTER  # type: ignore
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet  # type: ignore
from reportlab.lib.units import inch  # type: ignore
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer  # type: ignore


class DocumentGenerator:
    """Convert tailored resume / cover letter content into DOCX and PDF."""

    def generate_resume_docx(
        self,
        user_name: str,
        contact_info: str,
        summary: str,
        bullets: list[str],
        skills: Optional[list[str]] = None,
        base_excerpt: Optional[str] = None,
    ) -> io.BytesIO:
        doc = Document()

        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(user_name.upper())
        name_run.bold = True
        name_run.font.size = Pt(16)

        if contact_info:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_para.add_run(contact_info)

        if summary:
            doc.add_heading("Professional Summary", level=1)
            doc.add_paragraph(summary)

        if bullets:
            doc.add_heading("Key Achievements & Experience", level=1)
            for bullet in bullets:
                doc.add_paragraph(bullet, style="List Bullet")

        if skills:
            doc.add_heading("Highlighted Skills", level=1)
            doc.add_paragraph(", ".join(skills))

        if base_excerpt:
            doc.add_heading("Supporting Experience (from base resume)", level=1)
            for block in self._split_blocks(base_excerpt, limit=12):
                doc.add_paragraph(block)

        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        return stream

    # Backward-compatible alias used by /documents/export/docx
    def generate_docx(self, user_name: str, contact_info: str, summary: str, bullets: list[str]) -> io.BytesIO:
        return self.generate_resume_docx(user_name, contact_info, summary, bullets)

    def generate_cover_letter_docx(self, user_name: str, content: str) -> io.BytesIO:
        doc = Document()
        for para in self._paragraphs(content):
            doc.add_paragraph(para)
        if user_name and user_name.lower() not in content.lower():
            doc.add_paragraph("")
            doc.add_paragraph(user_name)
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        return stream

    def generate_resume_pdf(
        self,
        user_name: str,
        contact_info: str,
        summary: str,
        bullets: list[str],
        skills: Optional[list[str]] = None,
        base_excerpt: Optional[str] = None,
    ) -> io.BytesIO:
        """Prefer LaTeX (ATS-friendly single-column PDF); fall back to ReportLab."""
        try:
            return self._generate_resume_pdf_latex(
                user_name, contact_info, summary, bullets, skills, base_excerpt
            )
        except Exception as exc:
            print(f"[DocumentGenerator] LaTeX resume PDF failed ({exc}); using ReportLab fallback")
            return self._generate_resume_pdf_reportlab(
                user_name, contact_info, summary, bullets, skills, base_excerpt
            )

    def _generate_resume_pdf_latex(
        self,
        user_name: str,
        contact_info: str,
        summary: str,
        bullets: list[str],
        skills: Optional[list[str]] = None,
        base_excerpt: Optional[str] = None,
    ) -> io.BytesIO:
        import os
        import shutil
        import subprocess
        import tempfile

        from jinja2 import Template

        if not shutil.which("pdflatex"):
            raise RuntimeError("pdflatex not found on PATH")

        template_path = os.path.join(
            os.path.dirname(__file__), "..", "templates", "resume_template.tex"
        )
        with open(template_path, "r", encoding="utf-8") as f:
            template_str = f.read()

        template = Template(template_str)
        tex_content = template.render(
            user_name=self._escape_latex(user_name),
            contact_info=self._escape_latex(contact_info or ""),
            summary=self._escape_latex(summary or ""),
            bullets=[self._escape_latex(b) for b in (bullets or [])],
            skills=self._escape_latex(", ".join(skills) if skills else ""),
            base_excerpt_blocks=[
                self._escape_latex(b) for b in self._split_blocks(base_excerpt, limit=10)
            ]
            if base_excerpt
            else [],
        )

        with tempfile.TemporaryDirectory() as temp_dir:
            tex_file = os.path.join(temp_dir, "resume.tex")
            with open(tex_file, "w", encoding="utf-8") as f:
                f.write(tex_content)

            # Two passes help hyperref/refs settle for cleaner ATS-parseable PDFs
            for _ in range(2):
                result = subprocess.run(
                    ["pdflatex", "-interaction=nonstopmode", "-halt-on-error", "resume.tex"],
                    cwd=temp_dir,
                    check=False,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                )
                if result.returncode != 0:
                    err = (result.stderr or result.stdout or b"").decode("utf-8", errors="ignore")
                    raise RuntimeError(f"pdflatex failed: {err[-2000:]}")

            pdf_file = os.path.join(temp_dir, "resume.pdf")
            if not os.path.exists(pdf_file):
                raise RuntimeError("pdflatex completed but resume.pdf was not produced")
            with open(pdf_file, "rb") as f:
                pdf_data = f.read()

        stream = io.BytesIO(pdf_data)
        stream.seek(0)
        return stream

    def _generate_resume_pdf_reportlab(
        self,
        user_name: str,
        contact_info: str,
        summary: str,
        bullets: list[str],
        skills: Optional[list[str]] = None,
        base_excerpt: Optional[str] = None,
    ) -> io.BytesIO:
        stream = io.BytesIO()
        doc = SimpleDocTemplate(
            stream,
            pagesize=LETTER,
            leftMargin=0.75 * inch,
            rightMargin=0.75 * inch,
            topMargin=0.7 * inch,
            bottomMargin=0.7 * inch,
        )
        styles = getSampleStyleSheet()
        title = ParagraphStyle(
            "ResumeTitle",
            parent=styles["Heading1"],
            fontSize=16,
            alignment=1,
            spaceAfter=6,
        )
        contact = ParagraphStyle(
            "ResumeContact",
            parent=styles["Normal"],
            fontSize=9,
            alignment=1,
            spaceAfter=12,
        )
        heading = ParagraphStyle(
            "ResumeH",
            parent=styles["Heading2"],
            fontSize=12,
            spaceBefore=10,
            spaceAfter=6,
        )
        body = ParagraphStyle("ResumeBody", parent=styles["Normal"], fontSize=10, leading=13, spaceAfter=6)

        story: List[Any] = [Paragraph(self._escape(user_name.upper()), title)]
        if contact_info:
            story.append(Paragraph(self._escape(contact_info), contact))
        if summary:
            story.append(Paragraph("Professional Summary", heading))
            story.append(Paragraph(self._escape(summary), body))
        if bullets:
            story.append(Paragraph("Experience", heading))
            items = [
                ListItem(Paragraph(self._escape(b), body), leftIndent=12, bulletColor="black")
                for b in bullets
                if b
            ]
            if items:
                story.append(ListFlowable(items, bulletType="bullet", start="•"))
        if skills:
            story.append(Paragraph("Skills", heading))
            story.append(Paragraph(self._escape(", ".join(skills)), body))
        if base_excerpt:
            story.append(Paragraph("Additional Experience", heading))
            for block in self._split_blocks(base_excerpt, limit=12):
                story.append(Paragraph(self._escape(block), body))

        doc.build(story)
        stream.seek(0)
        return stream


    def generate_cover_letter_pdf(self, user_name: str, content: str) -> io.BytesIO:
        stream = io.BytesIO()
        doc = SimpleDocTemplate(
            stream,
            pagesize=LETTER,
            leftMargin=0.85 * inch,
            rightMargin=0.85 * inch,
            topMargin=0.85 * inch,
            bottomMargin=0.85 * inch,
        )
        styles = getSampleStyleSheet()
        body = ParagraphStyle("CLBody", parent=styles["Normal"], fontSize=11, leading=15, spaceAfter=10)
        story: List[Any] = []
        for p in self._paragraphs(content):
            story.append(Paragraph(self._escape(p), body))
            
        if user_name and user_name.lower() not in content.lower():
            story.append(Spacer(1, 12))
            story.append(Paragraph(self._escape(user_name), body))
        doc.build(story)
        stream.seek(0)
        return stream

    @staticmethod
    def _escape(text: str) -> str:
        return (
            (text or "")
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

    @staticmethod
    def _escape_latex(text: str) -> str:
        if not text:
            return ""
        escaped = text.replace('\\', '\\textbackslash{}')
        escaped = escaped.replace('{', '\\{').replace('}', '\\}')
        escaped = escaped.replace('$', '\\$').replace('&', '\\&')
        escaped = escaped.replace('%', '\\%').replace('#', '\\#')
        escaped = escaped.replace('_', '\\_')
        escaped = escaped.replace('~', '\\textasciitilde{}').replace('^', '\\textasciicircum{}')
        return escaped

    @staticmethod
    def _paragraphs(content: str) -> List[str]:
        parts = re.split(r"\n\s*\n", (content or "").strip())
        out = []
        for part in parts:
            cleaned = " ".join(line.strip() for line in part.splitlines() if line.strip())
            if cleaned:
                out.append(cleaned)
        return out or [content or ""]

    @staticmethod
    def _split_blocks(text: str, limit: int = 12) -> List[str]:
        lines = [ln.strip(" •-\t") for ln in (text or "").splitlines() if ln.strip()]
        # Prefer bullet-like lines
        bullets = [ln for ln in lines if len(ln) > 40][:limit]
        return bullets or lines[:limit]
